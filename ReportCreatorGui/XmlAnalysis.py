# -*- coding: utf-8 -*-
import os.path
import sys
# import wx

from datetime import datetime
from docx import Document, opc
from docx.shared import Pt, RGBColor
from NiXml import StimulusProfile, Testrun, ErrorElements
import XmlParser


class ReportGenerate:
    def __init__(self, xml_file, ms_word_file):
        """
        creates a MS Word Report from XML File.

        :param xml_file:
        :param ms_word_file:
        """
        self.xml_file = xml_file
        self.ms_word_file_template = "Report_Template.docx"
        self.xml_parser = XmlParser.XmlParser(self.xml_file)

        self.testrun_l = []
        self.createTestruns()
        self.errors_l = []

        ## statistic variables ##
        self.nTestruns = len(self.testrun_l)

        if self.nTestruns == 0:
            # TODO: Outputdialog
            print("Keine Testruns in XML File gefunden.")
            print("Programm wird beendet.")
            #sys.exit()

        self.count_ok_testruns = 0
        self.count_nok_testruns = 0
        self.count_ntestruns = 0
        self.runtime_testruns = 0

        self.analyse_testruns()
        self.StimulusProfileInfo = self.createStimulusProfile()
        self.analyse_errors()


        # self.create_ms_word_report()

    def analyse_errors(self):
        """
        create error elements

        :return:
        """

        for error_el in self.xml_parser.get_error_elements():
            self.errors_l.append(ErrorElements(error_el))

    def analyse_testruns(self):
        """
        count ok/nok in testruns and summarise runningtime of each testrun
        :return: None
        """

        for test in self.testrun_l:
            if "Passed" in test.result:
                self.count_ok_testruns = self.count_ok_testruns + 1
            else:
                self.count_nok_testruns = self.count_nok_testruns + 1

            self.count_ntestruns = self.count_ntestruns + 1
            self.runtime_testruns = self.runtime_testruns + test.runningTime_seconds

    def createTestruns(self):
        """
        create Testrun xml tree objects
        :return: None
        """
        for el in self.xml_parser.get_testruns():
            self.testrun_l.append(Testrun(el))

    def createStimulusProfile(self):
        """
        create StimulusProfile xml tree objects
        :return: None
        """
        return StimulusProfile(self.xml_parser.get_root_tree())

    def create_ms_word_report(self, outputDir):
        """
        Creates the microsoft word report from template

        :return: None
        """
        try:
            document = Document(self.ms_word_file_template)
            paragraph = document.add_paragraph()

            # print sections to document
            self.print_section_StimulusProfile(paragraph)

            self.print_section_testrun(paragraph)

            self.print_section_statistics(paragraph)

            if len(self.errors_l):
                self.print_section_error(paragraph)

            date = datetime.now().isoformat()
            date = date.replace(":", "_")
            document.save(os.path.join(outputDir, "Report_Date_" + date + ".docx"))

        except opc.exceptions.PackageNotFoundError:
            # TODO: call dialog
            print("Word Template not found" + self.ms_word_file)
            raise

        except IOError:
            # TODO: call dialog
            print("Report konnte nicht erstellt werden.")
            raise

        except Exception, e:
            print("Report Erstellung fehlgeschlagen")
            raise

    def print_section_StimulusProfile(self, paragraph):
        """
        print stimulus profile information

        :param paragraph: python docx paragraph
        :return:
        """
        self.add_formated_run(paragraph, "Stimulus Profile:\n\n", 11, True)
        self.add_formated_run(paragraph, "Name:\t", 9.5, True)
        self.add_formated_run(paragraph, self.StimulusProfileInfo.name + "\n", 9.5)
        self.add_formated_run(paragraph, "Pfad:\t", 9.5, True)
        self.add_formated_run(paragraph, self.StimulusProfileInfo.path + "\n", 9.5)

        # Wenn Beschreibung vorhanden
        if self.StimulusProfileInfo.description:
            self.add_formated_run(paragraph, "Beschreibung:\t", 9.5, True)
            self.add_formated_run(paragraph, self.StimulusProfileInfo.description + "\n\n", 9.5)
        else:
            self.add_newline(paragraph, 2)

    def print_section_testrun(self, paragraph):
        """
        print all Testruns to document

        :param paragraph: python docx paragraph
        :return:
        """

        self.add_formated_run(paragraph, "Testrun:\n", 11, True)
        self.draw_hline(paragraph, 95)
        self.add_newline(paragraph, 1)

        for i, testruns in enumerate(self.testrun_l):
            i = i + 1
            self.add_formated_run(paragraph, "Nr." + str(i) + "\n", 11, True)
            self.add_formated_run(paragraph, "Name:\t", 11, True)
            self.add_formated_run(paragraph, testruns.name + "\n", 9.5)

            self.add_formated_run(paragraph, "Pfad:\t", 11, True)
            self.add_formated_run(paragraph, testruns.path + "\n", 9.5)

            # Wenn Beschreibung vorhanden
            if testruns.description:
                self.add_formated_run(paragraph, "Beschreibung:\t", 9.5, True)
                self.add_formated_run(paragraph, testruns.description + "\n\n", 9.5)
            else:
                self.add_newline(paragraph, 2)

            self.add_formated_run(paragraph, "Startzeit:\t", 11, True)
            self.add_formated_run(paragraph, testruns.starttime.strftime("Datum %d.%m.%Y %H:%M:%S") + "\n", 9.5,
                                  color=[165, 42, 42])

            self.add_formated_run(paragraph, "Endzeit:\t", 11, True)
            self.add_formated_run(paragraph, testruns.endtime.strftime("Datum %d.%m.%Y %H:%M:%S") + "\n", 9.5,
                                  color=[165, 42, 42])

            self.add_formated_run(paragraph, "Laufzeit:\t", 11, True)
            self.add_formated_run(paragraph, str(testruns.timedelta.total_seconds()) + " Sekunden\n", 9.5,
                                  color=[165, 42, 42])

            self.add_formated_run(paragraph, "Ergebnis:\t", 11, True)

            if "Passed" in testruns.result:
                self.add_formated_run(paragraph, "OK", 9.5, color=[0, 176, 80])
            else:
                self.add_formated_run(paragraph, "NOK", 9.5, color=[255, 0, 0])

            self.add_newline(paragraph, 1)

            self.draw_hline(paragraph, 95)
            self.add_newline(paragraph, 2)

    def print_section_error(self, paragraph):
        """
        print all errors

        :param paragraph: Pyth
        :return: None
        """
        self.add_newline(paragraph, 5)
        self.add_formated_run(paragraph, "Error:\n", 11, True)
        self.draw_hline(paragraph, 95)
        self.add_newline(paragraph, 1)

        for i, errors in enumerate(self.errors_l):
            i = i + 1
            self.add_formated_run(paragraph, "Nr." + str(i) + "\n", 11, True)
            self.add_formated_run(paragraph, "Name:\t", 11, True)
            self.add_formated_run(paragraph, errors.name + "\n", 9.5)

            self.add_formated_run(paragraph, "ID:\t", 11, True)
            self.add_formated_run(paragraph, errors.ID + "\n", 9.5)

            self.add_formated_run(paragraph, "Beschreibung:\t", 11, True)
            self.add_formated_run(paragraph, errors.ErrorDescription + "\n", 9.5)

            self.draw_hline(paragraph, 95)

    def print_section_statistics(self, paragraph):
        """
        :param paragraph:
        :return:
        """

        self.add_formated_run(paragraph, "Zusammenfassung:\n", 11, True)
        self.add_newline(paragraph, 1)

        self.add_formated_run(paragraph, "Testruns ", 11, bold_flag=True)
        self.add_formated_run(paragraph, "OK", 9.5, bold_flag=True, color=[0, 176, 80])
        self.add_formated_run(paragraph,
                              " :\t\t" + str(self.count_ok_testruns) + "/" + str(self.count_ntestruns) + "\n", 9.5)

        self.add_formated_run(paragraph, "Testruns ", 11, bold_flag=True)
        self.add_formated_run(paragraph, "NOK", 9.5, bold_flag=True, color=[255, 0, 0])
        self.add_formated_run(paragraph, " :\t" + str(self.count_nok_testruns) + "/" + str(self.count_ntestruns) + "\n",
                              9.5)

        self.add_formated_run(paragraph, "Gesamtlaufzeit:\t", 11, True)
        self.add_formated_run(paragraph, str(self.runtime_testruns), 9.5)

    def add_formated_run(self, paragraph, text, size=10, bold_flag=False, color=None):
        """
        wrapper function for formatted printing.
        Print at the current cursor position

        :param paragraph: paragraph object
        :param text: text which will be printed to document
        :param size: fontsize
        :param bold_flag: True or False
        :param color: rgb colorcode as pyhton list [r,g,b]
        :return:
        """
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold_flag

        if color:
            run.font.color.rgb = RGBColor(*color)

    def draw_hline(self, paragraph, length):
        """
        draw a horizontal lines from underscores at the current cursor position

        :param paragraph: paragraph object
        :param length: length of summarised underscores
        :return:
        """
        underline = "_"
        for uline in range(length):
            underline = underline + "_"

        run = paragraph.add_run(underline)

    def add_newline(self, paragraph, n):
        """
        add a newline at the current cursor position
        :param paragraph: paragraph object
        :param n: number of newlines
        :return:
        """

        newlines = ""
        for el in range(n):
            newlines = newlines + "\n"

        run = paragraph.add_run(newlines)


# def Warn(parent, message, caption='Warning!'):
#     dlg = wx.MessageDialog(parent, message, caption, wx.OK | wx.ICON_EXCLAMATION)
#     dlg.ShowModal()
#     dlg.Destroy()


def XmlAnalysis(xmlFile, outputDir):
    if not os.path.isfile(xmlFile):
        print("FEHLER: Konnte XML Datei nicht finden->" + xmlFile)
        print("Programm ende")
        sys.exit(1)

    if not os.path.isdir(outputDir):
        print("FEHLER: Ziel Ordner für Report nicht gefunden -> " + outputDir)
        print("Programm ende")
        sys.exit(1)

    RG = ReportGenerate(xmlFile, "Report_Template.docx")
    RG.create_ms_word_report(outputDir)
#    sys.exit(0)  # succes

